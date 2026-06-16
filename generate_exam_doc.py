from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

# ========== 辅助函数 ==========
def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.3)
    return p

def add_number(text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.3)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()  # spacing

# ========== 正文 ==========
add_title("软件开发方法学 — 期末考试复习资料")

# ==================== 第一章 ====================
add_h1("一、软件生命周期（Software Life Cycle）")

add_h2("1.1 定义")
aadd_para('软件生命周期是指软件产品从概念提出、开发、投入使用，直到最终退役（废弃）的整个过程。它涵盖了软件从\u201c诞生\u201d到\u201c死亡\u201d的所有阶段。')

add_h2("1.2 经典阶段划分")
add_table(
    ["阶段", "主要活动", "主要产出"],
    [
        ["可行性研究", "技术/经济/操作可行性分析", "可行性研究报告"],
        ["需求分析", "收集、分析、规格化用户需求", "需求规格说明书(SRS)"],
        ["概要设计（系统设计）", "架构设计、模块划分、接口定义", "概要设计文档"],
        ["详细设计", "模块内部算法与数据结构设计", "详细设计文档"],
        ["编码（实现）", "用编程语言编写代码", "源代码、可执行程序"],
        ["测试", "单元测试、集成测试、系统测试、验收测试", "测试报告"],
        ["部署与维护", "安装部署、纠错/适应/完善/预防性维护", "维护记录"],
        ["退役", "系统下线、数据迁移", "退役报告"],
    ],
)

add_h2("1.3 常见考题")
add_bullet("简答题：请描述软件生命周期的各个阶段及其主要任务。")
add_bullet("选择题：软件维护属于软件生命周期的哪个阶段？（最长、成本最高的阶段）")
add_bullet("论述题：为什么说软件维护是生命周期中最昂贵的阶段？如何降低维护成本？")

# ==================== 第二章 ====================
add_h1("二、软件开发过程模型")

add_h2("2.1 瀑布模型（Waterfall Model）")
add_h3("基本思想")
aadd_para('将软件开发过程划分为若干顺序相连的阶段，每个阶段完成后必须经过评审，才能进入下一阶段。如同瀑布流水，逐级下落，不可逆转。')
add_h3("阶段流程")
add_para("可行性研究 → 需求分析 → 概要设计 → 详细设计 → 编码 → 测试 → 维护")
add_h3("优点")
add_bullet("阶段清晰，文档规范，易于管理和控制")
add_bullet("适合需求明确、变化少的项目")
add_h3("缺点")
add_bullet("需求变更代价大，灵活性差")
add_bullet("用户要到项目后期才能看到产品，风险高")
add_bullet("各阶段之间信息反馈慢（"阻塞"现象）")
add_h3("适用场景")
aadd_para('需求明确且稳定的系统，如嵌入式系统、银行核心系统等。')

add_h2("2.2 螺旋模型（Spiral Model）")
add_h3("基本思想")
aadd_para('由 Barry Boehm 于 1986 年提出。将瀑布模型与快速原型模型相结合，强调风险分析。每一轮螺旋包括四个象限。')
add_h3("四个象限（每轮迭代）")
add_number("制定计划：确定目标、方案和约束条件")
add_number("风险分析：识别并分析技术和项目风险")
add_number("工程实施：开发、测试（可能包含原型）")
add_number("客户评估：用户评审，提出改进意见")
add_h3("优点")
add_bullet("强调风险分析，适合大型复杂项目")
add_bullet("迭代式开发，逐步完善")
add_bullet("用户参与度高")
add_h3("缺点")
add_bullet("风险分析需要高度专业知识，成本高")
add_bullet("不适合小型项目")
add_bullet("过程复杂，管理难度大")
add_h3("适用场景")
aadd_para('大型、高风险、需求不确定的复杂系统，如航天系统、军事系统。')

add_h2("2.3 喷泉模型（Fountain Model）")
add_h3("基本思想")
aadd_para('喷泉模型是一种面向对象的软件开发过程模型。强调开发过程的迭代性和无间隙性（各阶段之间没有明显边界），如同喷泉的水流上下往复。')
add_h3("核心特征")
add_bullet("迭代（Iteration）：分析、设计、编码等活动可以反复进行")
add_bullet("无间隙（Seamless）：各阶段之间没有严格的先后顺序，可以重叠")
add_bullet("以对象为核心驱动：整个开发过程围绕对象展开")
add_h3("优点")
add_bullet("支持面向对象方法，灵活性强")
add_bullet("提高开发效率，减少返工")
add_h3("缺点")
add_bullet("管理困难，需要良好的配置管理")
add_bullet("对开发人员的面向对象技术要求高")
add_h3("适用场景")
aadd_para('面向对象的软件开发项目。')

add_h2("2.4 三种模型对比")
add_table(
    ["对比维度", "瀑布模型", "螺旋模型", "喷泉模型"],
    [
        ["驱动方式", "文档驱动", "风险驱动", "对象驱动"],
        ["迭代性", "无（线性）", "有（螺旋迭代）", "有（往复迭代）"],
        ["风险分析", "无", "核心环节", "不突出"],
        ["阶段边界", "严格分明", "按象限划分", "无间隙"],
        ["用户参与", "低", "高", "中等"],
        ["适用规模", "中小型、需求明确", "大型、高风险", "面向对象项目"],
        ["方法论", "结构化方法", "综合方法", "面向对象方法"],
    ],
)

add_h2("2.5 常见考题")
add_bullet("简答题：比较瀑布模型和螺旋模型的区别。")
add_bullet("简答题：喷泉模型的"无间隙"和"迭代"分别指什么？")
add_bullet("选择题：以下哪种模型最强调风险分析？（螺旋模型）")
add_bullet("应用题：某银行要开发新的核心交易系统，需求明确，你会选择哪种过程模型？为什么？")
add_bullet("论述题：分析螺旋模型四个象限各自的作用。")

# ==================== 第三章 ====================
add_h1("三、功能测试")

add_h2("3.1 软件测试概述")
aadd_para('软件测试是在规定条件下对软件产品进行操作，以发现缺陷、验证是否满足需求的过程。测试分为功能测试和非功能测试（性能、安全等）。')

add_h2("3.2 单元测试（Unit Testing）")
add_h3("定义")
aadd_para('对软件中最小的可测试单元（通常是一个函数/方法/类）进行测试，验证其行为是否符合预期。')
add_h3("特点")
add_bullet("由开发人员编写和执行")
add_bullet("通常在编码阶段同步进行")
add_bullet("使用测试框架（如 JUnit、pytest、unittest）")
add_bullet("常使用 Mock/Stub 隔离外部依赖")
add_h3("常用技术")
add_bullet("等价类划分：将输入划分为有效等价类和无效等价类")
add_bullet("边界值分析：测试输入范围的边界值")
add_bullet("白盒测试（了解内部逻辑）：语句覆盖、分支覆盖、路径覆盖")
add_h3("示例")
aadd_para('对 add(a, b) 函数编写测试用例：正常输入 add(1,2)=3；边界输入 add(0,0)=0；异常输入 add(-1,1)=0 等。')

add_h2("3.3 集成测试（Integration Testing）")
add_h3("定义")
aadd_para('在单元测试基础上，将多个模块/组件组合在一起，测试它们之间的交互是否正确。')
add_h3("集成策略")
add_table(
    ["策略", "说明", "优缺点"],
    [
        ["自顶向下", "从顶层模块开始，逐步集成底层模块，用桩(Stub)代替未集成的模块", "可早期验证主控制逻辑；但需要大量桩模块"],
        ["自底向上", "从底层模块开始，逐步向上集成，用驱动(Driver)调用上层接口", "不需要桩模块；但顶层模块最后才能测试"],
        ["大爆炸", "所有模块一次性集成", "简单直接；但问题定位困难"],
        ["三明治/混合", "结合自顶向下和自底向上", "灵活性高，但管理复杂"],
    ],
)

add_h2("3.4 单元测试 vs 集成测试 对比")
add_table(
    ["对比维度", "单元测试", "集成测试"],
    [
        ["测试对象", "单个函数/方法/类", "多个模块之间的交互"],
        ["测试目的", "验证单元内部逻辑正确性", "验证模块间接口与协作正确性"],
        ["执行者", "开发人员", "开发人员或测试团队"],
        ["测试方法", "白盒为主", "灰盒/黑盒为主"],
        ["依赖处理", "使用 Mock/Stub 隔离", "使用真实组件或部分真实组件"],
        ["执行速度", "快", "较慢"],
        ["发现问题", "单元内部缺陷", "接口不匹配、数据传递错误等"],
    ],
)

add_h2("3.5 测试的V模型")
aadd_para('V模型将测试与开发阶段对应起来：')
add_bullet("需求分析 ←→ 验收测试（Acceptance Testing）")
add_bullet("概要设计 ←→ 系统测试（System Testing）")
add_bullet("详细设计 ←→ 集成测试（Integration Testing）")
add_bullet("编码 ←→ 单元测试（Unit Testing）")

add_h2("3.6 常见考题")
add_bullet("简答题：单元测试和集成测试的区别是什么？")
add_bullet("简答题：简述自顶向下和自底向上集成策略。")
add_bullet("设计题：给定一个用户登录模块，设计单元测试用例。")
add_bullet("选择题：集成测试中，桩模块(Stub)的作用是？（代替尚未集成的被调用模块）")

# ==================== 第四章 ====================
add_h1("四、数据流图（Data Flow Diagram, DFD）")

add_h2("4.1 定义")
aadd_para('数据流图是一种结构化分析工具，用于描述系统中数据的流动、处理和存储。它从数据传递和加工的角度，刻画系统的逻辑功能，不涉及物理实现。')

add_h2("4.2 四种基本元素")
add_table(
    ["符号", "名称", "含义"],
    [
        ["○（圆形/圆角矩形）", "加工（Process）", "对数据进行处理/变换"],
        ["→（箭头）", "数据流（Data Flow）", "数据的流动方向"],
        ["☰（开口矩形/双线）", "数据存储（Data Store）", "数据的静态存储（文件、数据库）"],
        ["□（矩形）", "外部实体（External Entity）", "系统外部的数据源或数据终点"],
    ],
)

add_h2("4.3 分层数据流图")
aadd_para('DFD通常采用自顶向下、逐层分解的方式绘制：')
aadd_bullet('顶层图（Context Diagram / Level-0 DFD）：将整个系统视为一个加工，展示系统与外部实体之间的数据流。')
add_bullet('0层图（Level-1 DFD）：将顶层加工分解为若干子加工。')
add_bullet('更底层图：继续分解，直到每个加工足够简单（基本加工）。')
add_para('原则：父图与子图必须保持平衡（输入输出数据流一致）。')

add_h2("4.4 绘制原则与常见错误")
add_bullet("守恒原则：加工的输入数据流应足以产生输出数据流")
add_bullet("平衡原则：子图的输入/输出必须与父图对应加工的输入/输出一致")
add_bullet("常见错误：黑洞（加工只有输入没有输出）、奇迹（加工只有输出没有输入）、灰洞（输入不足以产生输出）")

add_h2("4.5 数据字典")
aadd_para('数据字典是DFD的补充，用于定义DFD中所有数据元素和数据结构的详细信息，包括数据流、数据存储和数据项的组成。')

add_h2("4.6 常见考题")
add_bullet("绘图题：为"图书管理系统"画出顶层数据流图和0层数据流图。")
add_bullet("纠错题：指出以下DFD中的错误（黑洞/奇迹/灰洞/不平衡）。")
add_bullet("简答题：数据流图中的"黑洞"错误是什么？如何避免？")
add_bullet("简答题：顶层图和0层图之间的"平衡"原则是什么意思？")

# ==================== 第五章 ====================
add_h1("五、耦合与内聚")

add_h2("5.1 耦合（Coupling）")
add_para("耦合是指模块之间相互依赖的程度。耦合度越低，模块独立性越强，系统越容易维护和修改。设计目标是"低耦合"。")

add_h3("耦合类型（从低到高排列）")
add_table(
    ["耦合类型", "说明", "耦合度"],
    [
        ["无直接耦合", "两个模块之间没有直接联系", "最低 ★"],
        ["数据耦合", "模块间通过参数传递简单数据项", "低 ★★"],
        ["标记耦合（特征耦合）", "模块间传递整个数据结构，但只用其中部分字段", "中低 ★★★"],
        ["控制耦合", "一个模块传递控制信息（如标志位）给另一个模块", "中 ★★★★"],
        ["外部耦合", "多个模块共享同一全局变量或外部环境", "中高 ★★★★★"],
        ["公共耦合", "多个模块共享同一公共数据区（全局数据结构）", "高 ★★★★★★"],
        ["内容耦合", "一个模块直接访问/修改另一个模块的内部数据", "最高 ★★★★★★★"],
    ],
)

add_h2("5.2 内聚（Cohesion）")
add_para("内聚是指模块内部各元素（语句、函数）之间结合的紧密程度。内聚度越高，模块功能越单一、越专注。设计目标是"高内聚"。")

add_h3("内聚类型（从低到高排列）")
add_table(
    ["内聚类型", "说明", "内聚度"],
    [
        ["偶然内聚", "模块内的处理元素之间没有必然联系，只是偶然被放在一起", "最低 ★"],
        ["逻辑内聚", "模块内的处理在逻辑上属于同一类（如所有输入操作）", "低 ★★"],
        ["时间内聚", "模块内的处理在同一时间段内执行（如初始化）", "中低 ★★★"],
        ["过程内聚", "模块内的处理按特定顺序执行", "中 ★★★★"],
        ["通信内聚", "模块内的处理使用相同的数据或产生相同的数据", "中高 ★★★★★"],
        ["顺序内聚", "一个处理的输出是下一个处理的输入", "高 ★★★★★★"],
        ["功能内聚", "模块内所有处理共同完成一个且仅一个功能", "最高 ★★★★★★★"],
    ],
)

add_h2("5.3 设计原则总结")
add_para("高内聚低耦合（High Cohesion, Low Coupling）", bold=True)
add_bullet("高内聚：每个模块只负责一个明确的功能")
add_bullet("低耦合：模块之间通过简单的接口通信，减少相互依赖")
add_bullet("好处：提高可维护性、可复用性、可测试性")

add_h2("5.4 常见考题")
add_bullet("选择题：以下哪种耦合的耦合度最低？（数据耦合）")
add_bullet("选择题：以下哪种内聚的内聚度最高？（功能内聚）")
add_bullet("简答题：解释"高内聚低耦合"的设计原则及其优点。")
add_bullet("应用题：给定一段代码，判断模块间的耦合类型和模块的内聚类型，并提出改进方案。")
add_bullet("判断题：控制耦合比数据耦合的耦合度更高。（正确）")

# ==================== 第六章 ====================
add_h1("六、架构模式")

add_h2("6.1 MVC 模式（Model-View-Controller）")
add_h3("概述")
add_para("MVC 是一种将应用程序分为三个核心组件的架构模式，实现关注点分离（Separation of Concerns）。")

add_h3("三个组件")
add_table(
    ["组件", "职责", "举例"],
    [
        ["Model（模型）", "管理数据和业务逻辑，独立于UI", "数据库操作、业务规则类"],
        ["View（视图）", "负责数据的展示（用户界面）", "HTML页面、UI组件"],
        ["Controller（控制器）", "接收用户输入，调用模型，选择视图", "路由处理、请求调度"],
    ],
)

add_h3("工作流程")
add_number("用户通过 View 发起操作（如点击按钮）")
add_number("Controller 接收请求，调用 Model 处理业务逻辑")
add_number("Model 更新数据，通知 View")
add_number("View 从 Model 获取最新数据并展示")

add_h3("优点")
add_bullet("关注点分离，各组件独立开发和测试")
add_bullet("可维护性和可扩展性好")
add_bullet("支持多个 View 共享同一个 Model")
add_h3("缺点")
add_bullet("增加了系统复杂度")
add_bullet("对于简单应用可能过度设计")

add_h2("6.2 管道-过滤器模式（Pipe and Filter）")
add_h3("概述")
add_para("管道-过滤器模式将数据处理过程分解为一系列独立的处理步骤（过滤器），通过管道（Pipe）连接，数据以流的方式在过滤器之间传递。")

add_h3("核心概念")
add_bullet("过滤器（Filter）：完成独立的数据处理/变换功能。每个过滤器有输入端口和输出端口。")
add_bullet("管道（Pipe）：连接过滤器的输出端口和下一个过滤器的输入端口，负责数据传输。")

add_h3("特点")
add_bullet("过滤器之间相互独立，不知道上下游是谁")
add_bullet("支持过滤器的复用和重新组合")
add_bullet("支持并行执行（不同过滤器可以同时处理不同数据段）")
add_bullet("易于添加、删除、替换过滤器")

add_h3("优点")
add_bullet("组件可复用")
add_bullet("支持并行处理")
add_bullet("易于维护和扩展")
add_bullet("支持快速原型（可以灵活组合过滤器）")
add_h3("缺点")
add_bullet("不适合交互式系统")
add_bullet("数据格式转换开销")
add_bullet("过滤器之间的共享状态管理困难")

add_h3("典型应用")
add_bullet("编译器：词法分析 → 语法分析 → 语义分析 → 代码生成")
add_bullet("Unix Shell 管道命令：cat file | grep "error" | sort | uniq")
add_bullet("数据处理流水线（ETL）")

add_h2("6.3 MVC vs 管道-过滤器 对比")
add_table(
    ["对比维度", "MVC", "管道-过滤器"],
    [
        ["适用场景", "交互式应用（Web、GUI）", "数据流处理系统"],
        ["核心思想", "关注点分离（数据/展示/控制）", "数据流水线处理"],
        ["组件关系", "三者协作、相互引用", "线性串联、互不知晓"],
        ["数据流", "双向（用户交互）", "单向（数据流驱动）"],
        ["典型应用", "Web应用、桌面GUI", "编译器、Unix管道、ETL"],
    ],
)

add_h2("6.4 其他常见架构模式（扩展）")
add_table(
    ["模式", "核心思想"],
    [
        ["分层架构（Layered）", "将系统分为若干层，每层只为上一层提供服务（如表现层/业务层/数据层）"],
        ["客户端-服务器（C/S）", "客户端请求服务，服务器提供服务"],
        ["微服务架构", "将应用拆分为一组小型、自治的服务，每个服务独立部署"],
        ["事件驱动架构", "组件通过事件（发布/订阅）进行通信"],
    ],
)

add_h2("6.5 常见考题")
add_bullet("简答题：描述MVC模式中三个组件各自的职责。")
add_bullet("简答题：管道-过滤器模式的优缺点是什么？")
add_bullet("应用题：为一个在线商城系统设计MVC架构，说明各组件的职责。")
add_bullet("选择题：Unix Shell命令管道体现了哪种架构模式？（管道-过滤器）")
add_bullet("论述题：比较MVC模式和管道-过滤器模式，分别适用于什么场景？")

# ==================== 第七章 ====================
add_h1("七、补充知识点（可能考点）")

add_h2("7.1 软件工程的三要素")
add_bullet("方法（Method）：软件开发的技术方法（如面向对象方法、结构化方法）")
add_bullet("工具（Tool）：支持方法的软件工具（如IDE、CASE工具）")
add_bullet("过程（Process）：将方法和工具结合起来的框架（如RUP、敏捷）")

add_h2("7.2 其他开发过程模型")
add_table(
    ["模型", "核心思想", "优缺点"],
    [
        ["增量模型", "分批交付，每批增加功能", "用户可早期使用核心功能；但需要良好的模块划分"],
        ["演化模型/原型模型", "先快速构建原型，用户反馈后迭代改进", "适合需求不明确的项目；但原型可能被滥用"],
        ["敏捷开发（Agile）", "短迭代（Sprint），快速交付，拥抱变化", "灵活适应需求变化；但对团队协作要求高"],
        ["RAD（快速应用开发）", "利用组件和工具快速构建", "开发速度快；但适用范围有限"],
    ],
)

add_h2("7.3 软件质量保证")
add_bullet("验证（Verification）：软件是否正确实现了规格说明（"Are we building the product right?"）")
add_bullet("确认（Validation）：软件是否满足用户需求（"Are we building the right product?"）")
add_bullet("评审（Review）：包括走查（Walkthrough）和审查（Inspection）")

add_h2("7.4 UML 图（可能涉及）")
add_bullet("用例图（Use Case Diagram）：描述系统功能与参与者的关系")
add_bullet("类图（Class Diagram）：描述系统的静态结构")
add_bullet("序列图（Sequence Diagram）：描述对象间的交互顺序")
add_bullet("活动图（Activity Diagram）：描述业务流程")
add_bullet("状态图（State Diagram）：描述对象的状态变化")

# ==================== 附录 ====================
add_h1("附录：高频考题汇总")

add_h2("A. 简答题精选")
add_number("请简述软件生命周期的各个阶段。")
add_number("比较瀑布模型和螺旋模型的区别，各适用于什么场景？")
add_number("解释"高内聚低耦合"的含义，并说明为什么这是好的设计原则。")
add_number("简述单元测试和集成测试的区别。")
add_number("什么是数据流图？它有哪些基本元素？")
add_number("MVC模式中三个组件各自承担什么职责？")
add_number("管道-过滤器模式有什么优缺点？")
add_number("喷泉模型的核心特征是什么？")

add_h2("B. 选择题精选")
add_number("以下哪个不属于软件生命周期的阶段？ A.需求分析 B.编码 C.市场调研 D.维护")
add_number("螺旋模型最核心的环节是？ A.编码 B.风险分析 C.测试 D.文档")
add_number("以下哪种耦合度最低？ A.控制耦合 B.数据耦合 C.公共耦合 D.内容耦合")
add_number("以下哪种内聚度最高？ A.时间内聚 B.通信内聚 C.功能内聚 D.逻辑内聚")
add_number("集成测试中，自顶向下集成需要使用？ A.驱动模块 B.桩模块 C.两者都需要 D.两者都不需要")
add_number("数据流图中，加工只有输入没有输出，属于什么错误？ A.奇迹 B.灰洞 C.黑洞 D.不平衡")

add_h2("C. 答案参考")
add_bullet("B1: C（市场调研不属于软件生命周期阶段）")
add_bullet("B2: B（风险分析是螺旋模型的核心）")
add_bullet("B3: B（数据耦合的耦合度最低）")
add_bullet("B4: C（功能内聚的内聚度最高）")
add_bullet("B5: B（自顶向下需要使用桩模块代替底层未集成模块）")
add_bullet("B6: C（只有输入没有输出是黑洞错误）")

doc.save("/Users/nate/Documents/Projects/python-projects/trailer/软件开发方法学_复习资料.docx")
print("✅ 文档已生成：软件开发方法学_复习资料.docx")
